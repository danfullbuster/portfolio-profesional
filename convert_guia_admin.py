"""
Script para convertir documentacion_total.md a Word (.docx)
"""
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
    
    # Crear documento
    doc = Document()
    
    # Configurar estilos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Leer el archivo markdown
    with open('documentacion_total.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Título principal
    title = doc.add_heading('Documentación Completa - Portfolio MiCV', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Procesar línea por línea
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
        
        # Encabezados
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        
        # Listas
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
        
        # Bloques de código
        elif line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text)
            p.style = 'No Spacing'
            run = p.runs[0]
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        
        # Tablas (formato simple)
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            # Detectar tabla
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                if not lines[i].strip().replace('|', '').replace('-', '').strip():
                    i += 1
                    continue
                table_lines.append(lines[i].strip())
                i += 1
            
            if len(table_lines) > 1:
                # Crear tabla
                cols = len([c for c in table_lines[0].split('|') if c.strip()])
                rows = len(table_lines)
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_text in enumerate(table_lines):
                    cells = [c.strip() for c in row_text.split('|') if c.strip()]
                    for col_idx, cell_text in enumerate(cells):
                        if col_idx < cols:
                            table.rows[row_idx].cells[col_idx].text = cell_text
            i -= 1
        
        # Texto normal
        else:
            # Eliminar markdown de negrita y cursiva
            text = line.replace('**', '').replace('__', '').replace('*', '').replace('_', '')
            if text:
                doc.add_paragraph(text)
        
        i += 1
    
    # Guardar documento
    doc.save('documentacion_total.docx')
    print("Archivo documentacion_total.docx creado exitosamente!")
    
except ImportError:
    print("Instalando python-docx...")
    import subprocess
    subprocess.run(['pip', 'install', 'python-docx'], check=True)
    print("Instalado. Por favor ejecuta el script nuevamente.")
except Exception as e:
    print(f"Error: {e}")
